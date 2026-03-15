import pandas as pd
import numpy as np
import io
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session, joinedload
from typing import List
import os
import time
import shutil
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import FileResponse
from docxtpl import DocxTemplate

from app.api import deps
from app.models.user import User
from app.models.contratacion import TipoProceso, PlantillaDocumento, ProcesoContratacion, DocumentoGenerado, PacAnual, ItemPac, ProcesoItemPacLink, HistoricoReformaPac, GenealogiaMontoPac, StatusItemPac
from app.schemas.contratacion import (
    TipoProcesoCreate, TipoProcesoUpdate,
    TipoProcesoResponse, PlantillaDocumentoResponse, 
    ProcesoContratacionCreate, ProcesoContratacionResponse,
    GenerarDocumentoRequest, RegenerarDocumentoRequest, PacAnualCreate, PacAnualResponse, ItemPacCreate, ItemPacResponse, ProcesoItemPacLinkCreate, ReformaPacCreate, MovimientoReformaCreate, DocumentoGeneradoResponse
)

router = APIRouter()

def get_render_data(doc, req_datos):
    from docxtpl import InlineImage
    from docx.shared import Mm
    import base64
    import io

    def procesar_valor(val):
        # Si es un string base64 de imagen
        if isinstance(val, str) and val.startswith("data:image"):
            try:
                header, encoded = val.split(",", 1)
                image_data = base64.b64decode(encoded)
                image_stream = io.BytesIO(image_data)
                return InlineImage(doc, image_stream, width=Mm(50))
            except Exception:
                return val
        # Si es una lista (para tablas dinamicas), aplicar recursividad
        elif isinstance(val, list):
            nueva_lista = []
            for item in val:
                if isinstance(item, dict):
                    nueva_lista.append(procesar_valor(item))
                else:
                    nueva_lista.append(item)
            return nueva_lista
        # Si es un diccionario
        elif isinstance(val, dict):
            nuevo_dict = {}
            for k, v in val.items():
                nuevo_dict[k] = procesar_valor(v)
            return nuevo_dict
        
        # Cualquier otra cosa se queda igual
        return val

    return procesar_valor(req_datos)


# ----- PROCESOS DE CONTRATACIÓN -----
@router.get("/procesos", response_model=List[ProcesoContratacionResponse])
def get_procesos(db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    return db.query(ProcesoContratacion).all()

@router.post("/procesos", response_model=ProcesoContratacionResponse)
def create_proceso(item_in: ProcesoContratacionCreate, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    data = item_in.model_dump()
    items_pac = data.pop("items_pac", [])
    
    db_item = ProcesoContratacion(**data, usuario_id=current_user.id)
    db.add(db_item)
    db.flush() # get id
    
    for link in items_pac:
        db_link = ProcesoItemPacLink(
            proceso_id=db_item.id,
            item_pac_id=link['item_pac_id'],
            monto_comprometido=link['monto_comprometido']
        )
        db.add(db_link)
        
    db.commit()
    db.refresh(db_item)
    return db_item

@router.get("/procesos/{id}", response_model=ProcesoContratacionResponse)
def get_proceso(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    db_item = db.query(ProcesoContratacion).filter(ProcesoContratacion.id == id).first()
    if not db_item:
        raise HTTPException(404, "Proceso no encontrado")
    return db_item

# ----- PLANTILLAS -----

@router.get("/tipos", response_model=List[TipoProcesoResponse])
def get_tipos(db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    return db.query(TipoProceso).all()

@router.post("/tipos", response_model=TipoProcesoResponse)
def create_tipo(item_in: TipoProcesoCreate, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    db_item = TipoProceso(**item_in.model_dump(), creado_por_id=current_user.id)
    db.add(db_item)
    db.commit()
    db.refresh(db_item)
    return db_item

@router.put("/tipos/{id}", response_model=TipoProcesoResponse)
def update_tipo(id: int, item_in: TipoProcesoUpdate, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    db_item = db.query(TipoProceso).filter(TipoProceso.id == id).first()
    if not db_item:
        raise HTTPException(404, "Tipo de proceso no encontrado")
        
    update_data = item_in.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_item, field, value)
        
    db_item.actualizado_por_id = current_user.id
    db.commit()
    db.refresh(db_item)
    return db_item


@router.post("/plantillas/upload", response_model=PlantillaDocumentoResponse)
def upload_plantilla(
    nombre: str = Form(...),
    tipo_proceso_id: int = Form(...),
    anio: int = Form(2026),
    file: UploadFile = File(...),
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
):
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "El archivo debe ser .docx")
        
    old_tpl = db.query(PlantillaDocumento).filter_by(nombre=nombre, tipo_proceso_id=tipo_proceso_id, is_activa=True).first()
    new_version = 1
    
    if old_tpl:
        old_tpl.is_activa = False
        new_version = old_tpl.version + 1
        db.add(old_tpl)
        
    os.makedirs("templates/contratacion", exist_ok=True)
    safe_name = nombre.replace(' ', '_')
    file_path = f"templates/contratacion/{tipo_proceso_id}_{safe_name}_v{new_version}_{int(time.time())}.docx"
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    new_tpl = PlantillaDocumento(
        nombre=nombre, 
        ruta_archivo_docx=file_path, 
        tipo_proceso_id=tipo_proceso_id, 
        anio=anio, 
        version=new_version, 
        is_activa=True,
        creado_por_id=current_user.id
    )
    db.add(new_tpl)
    db.commit()
    db.refresh(new_tpl)
    return new_tpl

from typing import Optional
from fastapi import Query

@router.get("/plantillas", response_model=List[PlantillaDocumentoResponse])
def get_plantillas(
    tipo_proceso_id: Optional[int] = Query(None),
    is_activa: Optional[bool] = Query(None),
    db: Session = Depends(deps.get_db), 
    current_user: User = Depends(deps.get_current_user)
):
    query = db.query(PlantillaDocumento)
    if tipo_proceso_id is not None:
        query = query.filter(PlantillaDocumento.tipo_proceso_id == tipo_proceso_id)
    if is_activa is not None:
        query = query.filter(PlantillaDocumento.is_activa == is_activa)
    return query.all()


@router.get("/plantillas/{id}/esquema")
def get_plantilla_esquema(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    plantilla = db.query(PlantillaDocumento).filter(PlantillaDocumento.id == id).first()
    if not plantilla:
        raise HTTPException(404, "Plantilla no encontrada")
        
    if not os.path.exists(plantilla.ruta_archivo_docx):
        raise HTTPException(400, "El archivo fisico no existe")
        
    import docx
    import re
    try:
        doc_docx = docx.Document(plantilla.ruta_archivo_docx)
        texto_completo = ""
        
        for p in doc_docx.paragraphs:
            texto_completo += p.text + "\n"
        for table in doc_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_completo += cell.text + "\n"
        # Extraer texto de encabezados y pies de pagina en todas sus variantes
        for section in doc_docx.sections:
            elementos_hf = [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer
            ]
            for hf in elementos_hf:
                if hf:
                    for p in hf.paragraphs:
                        texto_completo += p.text + "\n"
                    for table in hf.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                texto_completo += cell.text + "\n"

        
        # 1. Buscar todos los bucles y sus iteradores
        patron_bucle = r'{%\s*(?:tr\s+|p\s+)?for\s+(\w+)\s+in\s+(\w+)\s*%}'
        bucles_encontrados = re.findall(patron_bucle, texto_completo)
        mapa_iteradores = {iterador: nombre_lista for iterador, nombre_lista in bucles_encontrados}
        
        # 2. Buscar todas las variables con finditer para preservar el orden y sacar contexto
        patron_vars = r'\{\{\s*([\w.]+)\s*\}\}'
        matches = list(re.finditer(patron_vars, texto_completo))
        
        esquema_dict = {} # Preserva orden de insercion en Python 3.7+
        tablas_dict = {} # Preserva sub_atributos ordenados
        
        for match in matches:
            var_raw = match.group(1)
            
            # Omitir variables internas de jinja2 (ej: loop.index)
            if 'loop' in var_raw.lower():
                continue
                
            # Extraer contexto (40 chars)
            start_pos = max(0, match.start() - 40)
            end_pos = min(len(texto_completo), match.end() + 40)
            contexto = "..." + texto_completo[start_pos:match.start()].replace('\n', ' ').strip() + f" [ {var_raw} ] " + texto_completo[match.end():end_pos].replace('\n', ' ').strip() + "..."
            
            # Si tiene punto, pertenece a un iterador de tabla
            if '.' in var_raw:
                partes = var_raw.split('.', 1)
                iterador = partes[0]
                atributo = partes[1]
                
                if iterador in mapa_iteradores:
                    nombre_lista = mapa_iteradores[iterador]
                    if nombre_lista not in tablas_dict:
                        tablas_dict[nombre_lista] = {}
                    if atributo not in tablas_dict[nombre_lista]:
                        tablas_dict[nombre_lista][atributo] = contexto
                else:
                    if var_raw not in esquema_dict:
                        esquema_dict[var_raw] = {"tipo": "texto", "contexto": contexto}
            else:
                if var_raw not in esquema_dict:
                    if var_raw.startswith("img_"):
                        esquema_dict[var_raw] = {"tipo": "imagen", "contexto": contexto}
                    else:
                        esquema_dict[var_raw] = {"tipo": "texto", "contexto": contexto}
                
        # 4. Construir esquema JSON final
        esquema_final = []
        for nombre_var, datos in esquema_dict.items():
            esquema_final.append({
                "nombre": nombre_var,
                "tipo": datos["tipo"],
                "contexto": datos["contexto"]
            })
                
        for nombre_lista, atributos_dict in tablas_dict.items():
            # El frontend frontend usaba sub_atributos, pero lo unificaremos como columnas
            esquema_final.append({
                "nombre": nombre_lista,
                "tipo": "tabla_dinamica",
                "columnas": list(atributos_dict.keys()),
                "contexto": f"Matriz dinamica ({len(atributos_dict)} columnas)"
            })
            
        return {"variables": esquema_final}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analizando documento: {str(e)}")

# ----- DOCUMENTOS GENERADOS -----
@router.post("/documento")
def generar_documento(req: GenerarDocumentoRequest, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    # 1. Obtener plantilla
    plantilla = db.query(PlantillaDocumento).filter(PlantillaDocumento.id == req.plantilla_id).first()
    if not plantilla:
        raise HTTPException(404, "Plantilla no encontrada")
    
    if not os.path.exists(plantilla.ruta_archivo_docx):
        raise HTTPException(400, "Archivo físico de plantilla no existe")

        # 2. Generar DOCX
    from docxtpl import DocxTemplate
    doc = DocxTemplate(plantilla.ruta_archivo_docx)
    
    # Process base64
    datos_limpios = get_render_data(doc, req.datos)
    
    # Render explicitly without any env argument
    doc.render(datos_limpios)
    
    os.makedirs("generated", exist_ok=True)
    out_path = f"generated/doc_proc_{req.proceso_contratacion_id}_tpl_{req.plantilla_id}_v1_{int(time.time())}.docx"
    doc.save(out_path)

    # 3. Guardar en DB
    doc_gen = DocumentoGenerado(
        proceso_contratacion_id=req.proceso_contratacion_id,
        plantilla_id=req.plantilla_id,
        version=1,
        datos_json=req.datos,
        ruta_archivo_generado=out_path
    )
    db.add(doc_gen)
    db.commit()
    db.refresh(doc_gen)
    # 4. Devolver el archivo fisico descargable
    from fastapi.responses import FileResponse
    return FileResponse(
        path=out_path,
        filename=f"Proceso_Generado_v1.docx",
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={"Content-Disposition": f"attachment; filename=Proceso_Generado_v1.docx"}
    )

@router.get("/documento/{id}/datos")
def get_datos_documento(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    doc = db.query(DocumentoGenerado).filter(DocumentoGenerado.id == id).first()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")
    return doc.datos_json


@router.get("/documento/{id}/descargar")
def descargar_documento(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    doc = db.query(DocumentoGenerado).filter(DocumentoGenerado.id == id).first()
    if not doc or not os.path.exists(doc.ruta_archivo_generado):
        raise HTTPException(404, "Documento fisico no encontrado")
        
    from fastapi.responses import FileResponse
    return FileResponse(
        path=doc.ruta_archivo_generado,
        filename=f"Documento_v{doc.version}.docx",
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@router.put("/documento/{id}/regenerar")
def regenerar_documento(id: int, req: RegenerarDocumentoRequest, PacAnualCreate, PacAnualResponse, ItemPacCreate, ItemPacResponse, ProcesoItemPacLinkCreate, ReformaPacCreate, MovimientoReformaCreate, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    doc = db.query(DocumentoGenerado).filter(DocumentoGenerado.id == id).first()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")
        
    plantilla = db.query(PlantillaDocumento).filter(PlantillaDocumento.id == doc.plantilla_id).first()
    
        # Render new docx
    from docxtpl import DocxTemplate
    docx_tpl = DocxTemplate(plantilla.ruta_archivo_docx)
    datos_limpios = get_render_data(docx_tpl, req.datos)
    
    docx_tpl.render(datos_limpios)
    
    new_version = doc.version + 1
    out_path = f"generated/doc_proc_{doc.proceso_contratacion_id}_tpl_{doc.plantilla_id}_v{new_version}_{int(time.time())}.docx"
    docx_tpl.save(out_path)
    
    # Update DB
    doc.datos_json = req.datos
    doc.version = new_version
    doc.ruta_archivo_generado = out_path
    
    db.commit()
    db.refresh(doc)
    return doc

# --- PAC Endpoints ---
@router.post("/pac", response_model=PacAnualResponse, status_code=status.HTTP_201_CREATED)
def crear_pac(req: PacAnualCreate, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    db_obj = PacAnual(**req.dict())
    db.add(db_obj)
    db.commit()
    db.refresh(db_obj)
    return db_obj

@router.get("/pac", response_model=List[PacAnualResponse])
def listar_pac(db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    return db.query(PacAnual).options(joinedload(PacAnual.items)).all()


@router.get("/pac/{id}/items", response_model=List[ItemPacResponse])
def listar_items_pac(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    items = db.query(ItemPac).filter(ItemPac.pac_anual_id == id).all()
    return items


@router.get("/pac/items/all", response_model=List[ItemPacResponse])
def explorar_todos_items(db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    # Trae todos los items sin importar a que PAC pertenecen, usando join para traer data del padre
    return db.query(ItemPac).options(joinedload(ItemPac.pac)).all()

@router.post("/pac/{id}/items", response_model=ItemPacResponse, status_code=status.HTTP_201_CREATED)
def agregar_item_pac(id: int, req: ItemPacCreate, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    db_obj = ItemPac(**req.dict(), pac_anual_id=id)
    db.add(db_obj)
    db.commit()
    db.refresh(db_obj)
    return db_obj



# --- PAC Carga Masiva ---
@router.post("/pac/{id}/importar")
async def importar_pac(id: int, file: UploadFile = File(...), db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    pac = db.query(PacAnual).filter(PacAnual.id == id).first()
    if not pac:
        raise HTTPException(status_code=404, detail="PAC no encontrado")
        
    ext = file.filename.split(".")[-1].lower()
    if ext not in ["csv", "xls", "xlsx"]:
        raise HTTPException(status_code=400, detail="Formato no soportado. Usa CSV o Excel.")
        
    contents = await file.read()
    
    try:
        if ext == "csv":
            try:
                df = pd.read_csv(io.BytesIO(contents), encoding='utf-8')
            except UnicodeDecodeError:
                # Failsafe para archivos generados por Excel/Windows antiguos
                df = pd.read_csv(io.BytesIO(contents), encoding='latin1')
        else:
            df = pd.read_excel(io.BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el archivo: {str(e)}")

    # 1. Limpiar los nombres de columnas originales del archivo (quitar espacios sobrantes)
    df.columns = df.columns.str.strip()
    
    # 2. Mapeo Inteligente (Diccionario Traductor)
    mapeo_columnas = {
        'Partida Pres.': 'partida_presupuestaria',
        'CPC': 'cpc',
        'Tipo de Compra': 'tipo_compra',
        'Procedimiento': 'procedimiento',
        'Descripción': 'descripcion',
        'Cant.': 'cantidad',
        'Costo U.': 'costo_unitario',
        'V. Total': 'valor_total'
    }
    
    # Aplicar traduccion
    df = df.rename(columns=mapeo_columnas)
    
    # 3. Verificacion de Columnas Obligatorias
    columnas_requeridas = ['partida_presupuestaria', 'descripcion']
    columnas_faltantes = [col for col in columnas_requeridas if col not in df.columns]
    
    if columnas_faltantes:
        raise HTTPException(
            status_code=400, 
            detail=f"El archivo no tiene el formato correcto del SERCOP. Faltan las columnas mapeadas: {', '.join(columnas_faltantes)}"
        )

    # 4. Prevenir que SQLAlchemy explote por los NaNs nativos de Pandas
    df = df.replace({np.nan: None})

    def clean_money(val):
        if val is None: return 0.0
        if isinstance(val, (int, float)): return float(val)
        val_str = str(val).replace('$', '').replace(',', '').strip()
        try: return float(val_str)
        except: return 0.0


    # Lógica Anti-Duplicados
    items_existentes = db.query(ItemPac).filter(ItemPac.pac_anual_id == id).all()
    set_existentes = set()
    for ie in items_existentes:
        key = (str(ie.partida_presupuestaria).strip(), str(ie.cpc).strip(), str(ie.descripcion).strip()[:100])
        set_existentes.add(key)

    items_to_insert = []
    ignorados = 0

    
    for index, row in df.iterrows():
        try:
            partida = str(row['partida_presupuestaria']) if row['partida_presupuestaria'] else "N/A"
            if partida == "N/A" or partida.strip() == "" or partida == "None": continue # Skip empty lines
            
            cpc_val = str(row['cpc']) if 'cpc' in df.columns and row['cpc'] else ""
            desc_val = str(row['descripcion']) if row['descripcion'] else "Sin descripcion"
            
            # Verificación en Set
            row_key = (partida.strip(), cpc_val.strip(), desc_val.strip()[:100])
            if row_key in set_existentes:
                ignorados += 1
                continue
                
            item = ItemPac(
                pac_anual_id=id,
                partida_presupuestaria=partida,
                cpc=cpc_val if cpc_val else None,
                tipo_compra=str(row['tipo_compra']) if 'tipo_compra' in df.columns and row['tipo_compra'] else "Bien/Servicio",
                procedimiento=str(row['procedimiento']) if 'procedimiento' in df.columns and row['procedimiento'] else None,
                descripcion=desc_val,
                cantidad=clean_money(row['cantidad']) if 'cantidad' in df.columns else 1.0,
                costo_unitario=clean_money(row['costo_unitario']) if 'costo_unitario' in df.columns else 0.0,
                valor_total=clean_money(row['valor_total']) if 'valor_total' in df.columns else 0.0,
                status=StatusItemPac.ACTIVO
            )
            items_to_insert.append(item)
        except Exception as e:
            db.rollback()
            raise HTTPException(status_code=400, detail=f"Error en la fila Excel {index + 2}: {str(e)}")
            
    try:
        db.add_all(items_to_insert)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Error transaccional al guardar en la base de datos.")
        
    return {"ok": True, "filas_procesadas": len(items_to_insert), "ignoradas_por_duplicado": ignorados}

@router.post("/pac/{id}/reforma")

def reformar_pac(id: int, req: ReformaPacCreate, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    pac = db.query(PacAnual).filter(PacAnual.id == id).first()
    if not pac:
        raise HTTPException(status_code=404, detail="PAC no encontrado")
        
    # 1. Validar montos de movimientos
    suma_origen = sum([m.monto_transferido for m in req.movimientos])
    suma_destino = sum([i.valor_total for i in req.nuevos_items])
    
    # Tolerancia de decimales (centavos)
    if abs(suma_origen - suma_destino) > 0.01:
        raise HTTPException(status_code=400, detail=f"Desbalance financiero. El origen cede ${suma_origen} pero el destino requiere ${suma_destino}.")

    # 2. Actualizar versión del PAC
    pac.version_reforma = req.numero_reforma
    
    # 3. Crear Histórico
    reforma = HistoricoReformaPac(
        pac_anual_id=id,
        numero_reforma=req.numero_reforma,
        resolucion_administrativa=req.resolucion_administrativa,
        descripcion_justificacion=req.descripcion_justificacion
    )
    db.add(reforma)
    db.flush() # Para obtener reforma.id
    
    # 4. Reducir/Eliminar items origen
    for mov in req.movimientos:
        item_orig = db.query(ItemPac).filter(ItemPac.id == mov.item_origen_id).first()
        if not item_orig: continue
        
        # Reducir saldo
        item_orig.valor_total -= mov.monto_transferido
        item_orig.cantidad = 1 # simplificado para la reforma
        item_orig.costo_unitario = item_orig.valor_total
        
        if item_orig.valor_total <= 0.01:
            item_orig.status = StatusItemPac.ELIMINADO_POR_REFORMA
            item_orig.valor_total = 0
        else:
            item_orig.status = StatusItemPac.MODIFICADO_POR_REFORMA

    # 5. Crear items destino e inyectar genealogía
    for nuevo_item_req in req.nuevos_items:
        nuevo_item = ItemPac(**nuevo_item_req.dict(), pac_anual_id=id, status=StatusItemPac.ACTIVO)
        db.add(nuevo_item)
        db.flush() # Obtener nuevo_item.id
        
        # Enlazar genealogía (simplificado: asigna proporcionalmente)
        # En un sistema real se mapea 1 a 1 en el payload, pero asumimos que el array movimientos 
        # dicta de dónde sale la plata para toda la bolsa de nuevos items.
        for mov in req.movimientos:
            if mov.monto_transferido > 0:
                # Registramos el nexo
                gen = GenealogiaMontoPac(
                    historico_reforma_id=reforma.id,
                    item_origen_id=mov.item_origen_id,
                    item_destino_id=nuevo_item.id,
                    monto_transferido=mov.monto_transferido # Idealmente prorrateado
                )
                db.add(gen)
                
    db.commit()
    return {"ok": True, "reforma_id": reforma.id, "mensaje": "Reforma financiera ejecutada con cuadre exacto."}

@router.post("/procesos/{id}/vincular_pac")
def vincular_pac_proceso(id: int, req_links: List[ProcesoItemPacLinkCreate], db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    proceso = db.query(ProcesoContratacion).filter(ProcesoContratacion.id == id).first()
    if not proceso:
        raise HTTPException(status_code=404, detail="Proceso no encontrado")
    
    # Clean existing links
    db.query(ProcesoItemPacLink).filter(ProcesoItemPacLink.proceso_id == id).delete()
    
    for link in req_links:
        db_link = ProcesoItemPacLink(
            proceso_id=id,
            item_pac_id=link.item_pac_id,
            monto_comprometido=link.monto_comprometido
        )
        db.add(db_link)
    db.commit()
    return {"ok": True}

