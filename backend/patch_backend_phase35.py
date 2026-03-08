import re

path = "/home/ayepez/.openclaw/workspace/sge/backend/app/api/routes/contratacion.py"
with open(path, "r") as f:
    content = f.read()

# 1. Extract endpoint patch (Preserve order and extract context)
start_str = '@router.get("/plantillas/{id}/esquema")'
end_str = 'def generar_documento'
start_idx = content.find(start_str)
end_idx = content.find(end_str)

if start_idx != -1 and end_idx != -1:
    new_esquema = r"""@router.get("/plantillas/{id}/esquema")
def get_plantilla_esquema(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    plantilla = db.query(PlantillaDocumento).filter(PlantillaDocumento.id == id).first()
    if not plantilla:
        raise HTTPException(404, "Plantilla no encontrada")
    if not os.path.exists(plantilla.ruta_archivo_docx):
        raise HTTPException(400, "El archivo fisico no existe")
    import docx
    import re
    try:
        doc_docx = docx.Document(plantilla.ruta_archivo_docx)
        texto_completo = ""
        for p in doc_docx.paragraphs:
            texto_completo += p.text + "\n"
        for table in doc_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_completo += cell.text + "\n"
                    
        # Extraer variables y bucles manteniendo el orden
        pattern = r'({{\s*(\w+)\s*}}|{%\s*(?:tr\s+|p\s+)?for\s+\w+\s+in\s+(\w+)\s*%})'
        matches = list(re.finditer(pattern, texto_completo))
        
        seen = set()
        esquema = []
        for match in matches:
            full_match = match.group(1)
            var_name = match.group(2) or match.group(3)
            
            if not var_name or var_name in seen:
                continue
            seen.add(var_name)
            
            start_pos = max(0, match.start() - 40)
            end_pos = min(len(texto_completo), match.end() + 40)
            contexto = "..." + texto_completo[start_pos:match.start()].replace('\n', ' ').strip() + f" [ {var_name} ] " + texto_completo[match.end():end_pos].replace('\n', ' ').strip() + "..."
            
            if var_name.startswith("img_"):
                tipo = "imagen"
            elif full_match.startswith("{%") or var_name.startswith("tbl_") or var_name.startswith("lista_"):
                tipo = "tabla_dinamica"
            else:
                tipo = "texto"
                
            esquema.append({"nombre": var_name, "tipo": tipo, "contexto": contexto})
                
        return {"variables": esquema}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analizando documento: {str(e)}")

# ----- DOCUMENTOS GENERADOS -----
"""
    content = content[:start_idx] + new_esquema + content[end_idx:]

# 2. Add Base64 Image Processing Helper to the module
img_helper = """
def get_render_data(doc, req_datos):
    datos_para_render = req_datos.copy()
    from docxtpl import InlineImage
    from docx.shared import Mm
    import base64
    import io
    for key, val in req_datos.items():
        if isinstance(val, str) and val.startswith("data:image"):
            try:
                header, encoded = val.split(",", 1)
                image_data = base64.b64decode(encoded)
                image_stream = io.BytesIO(image_data)
                datos_para_render[key] = InlineImage(doc, image_stream, width=Mm(60))
            except Exception as e:
                pass
    return datos_para_render
"""
content = content.replace("router = APIRouter()", "router = APIRouter()\n" + img_helper)

# 3. Patch generar and regenerar to use get_render_data
content = content.replace("doc.render(req.datos)", "doc.render(get_render_data(doc, req.datos))")
content = content.replace("docx_tpl.render(req.datos)", "docx_tpl.render(get_render_data(docx_tpl, req.datos))")

with open(path, "w") as f:
    f.write(content)
