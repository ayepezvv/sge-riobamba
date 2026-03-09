import re

path = "/home/ayepez/.openclaw/workspace/sge/backend/app/api/routes/contratacion.py"
with open(path, "r") as f:
    content = f.read()

# Restore the context and order extraction
new_logic = """@router.get("/plantillas/{id}/esquema")
def get_plantilla_esquema(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    plantilla = db.query(PlantillaDocumento).filter(PlantillaDocumento.id == id).first()
    if not plantilla:
        raise HTTPException(404, "Plantilla no encontrada")
        
    if not os.path.exists(plantilla.ruta_archivo_docx):
        raise HTTPException(400, "El archivo fisico no existe")
        
    import docx
    import re
    try:
        doc_docx = docx.Document(plantilla.ruta_archivo_docx)
        texto_completo = ""
        
        for p in doc_docx.paragraphs:
            texto_completo += p.text + "\\n"
            
        for table in doc_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_completo += cell.text + "\\n"
        
        # 1. Buscar todos los bucles y sus iteradores
        patron_bucle = r'{%\\s*(?:tr\\s+|p\\s+)?for\\s+(\\w+)\\s+in\\s+(\\w+)\\s*%}'
        bucles_encontrados = re.findall(patron_bucle, texto_completo)
        mapa_iteradores = {iterador: nombre_lista for iterador, nombre_lista in bucles_encontrados}
        
        # 2. Buscar todas las variables con finditer para preservar el orden y sacar contexto
        patron_vars = r'\\{\\{\\s*([\\w.]+)\\s*\\}\\}'
        matches = list(re.finditer(patron_vars, texto_completo))
        
        esquema_dict = {} # Preserva orden de insercion en Python 3.7+
        tablas_dict = {} # Preserva sub_atributos ordenados
        
        for match in matches:
            var_raw = match.group(1)
            
            # Extraer contexto (40 chars)
            start_pos = max(0, match.start() - 40)
            end_pos = min(len(texto_completo), match.end() + 40)
            contexto = "..." + texto_completo[start_pos:match.start()].replace('\\n', ' ').strip() + f" [ {var_raw} ] " + texto_completo[match.end():end_pos].replace('\\n', ' ').strip() + "..."
            
            # Si tiene punto, pertenece a un iterador de tabla
            if '.' in var_raw:
                partes = var_raw.split('.', 1)
                iterador = partes[0]
                atributo = partes[1]
                
                if iterador in mapa_iteradores:
                    nombre_lista = mapa_iteradores[iterador]
                    if nombre_lista not in tablas_dict:
                        tablas_dict[nombre_lista] = {}
                    if atributo not in tablas_dict[nombre_lista]:
                        tablas_dict[nombre_lista][atributo] = contexto
                else:
                    if var_raw not in esquema_dict:
                        esquema_dict[var_raw] = {"tipo": "texto", "contexto": contexto}
            else:
                if var_raw not in esquema_dict:
                    if var_raw.startswith("img_"):
                        esquema_dict[var_raw] = {"tipo": "imagen", "contexto": contexto}
                    else:
                        esquema_dict[var_raw] = {"tipo": "texto", "contexto": contexto}
                
        # 4. Construir esquema JSON final
        esquema_final = []
        for nombre_var, datos in esquema_dict.items():
            esquema_final.append({
                "nombre": nombre_var,
                "tipo": datos["tipo"],
                "contexto": datos["contexto"]
            })
                
        for nombre_lista, atributos_dict in tablas_dict.items():
            # El frontend frontend usaba sub_atributos, pero lo unificaremos como columnas
            esquema_final.append({
                "nombre": nombre_lista,
                "tipo": "tabla_dinamica",
                "columnas": list(atributos_dict.keys()),
                "contexto": f"Matriz dinamica ({len(atributos_dict)} columnas)"
            })
            
        return {"variables": esquema_final}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analizando documento: {str(e)}")"""

start_idx = content.find('@router.get("/plantillas/{id}/esquema")')
end_idx = content.find('# ----- DOCUMENTOS GENERADOS -----', start_idx)

if start_idx != -1 and end_idx != -1:
    content = content[:start_idx] + new_logic + "\n\n" + content[end_idx:]
    with open(path, "w") as f:
        f.write(content)
