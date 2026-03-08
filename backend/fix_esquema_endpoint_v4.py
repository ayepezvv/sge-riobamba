import re
path = "/home/ayepez/.openclaw/workspace/sge/backend/app/api/routes/contratacion.py"
with open(path, "r") as f:
    content = f.read()

# Replace the get_plantilla_esquema function with the tank strategy, escaping properly for python scripts
new_func = r"""@router.get("/plantillas/{id}/esquema")
def get_plantilla_esquema(id: int, db: Session = Depends(deps.get_db), current_user: User = Depends(deps.get_current_user)):
    plantilla = db.query(PlantillaDocumento).filter(PlantillaDocumento.id == id).first()
    if not plantilla:
        raise HTTPException(404, "Plantilla no encontrada")
        
    if not os.path.exists(plantilla.ruta_archivo_docx):
        raise HTTPException(400, "El archivo fisico no existe")
        
    import docx
    import re
    try:
        doc_docx = docx.Document(plantilla.ruta_archivo_docx)
        texto_completo = ""
        
        # Extraer texto de parrafos
        for p in doc_docx.paragraphs:
            texto_completo += p.text + "\n"
            
        # Extraer texto de tablas
        for table in doc_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_completo += cell.text + "\n"
                    
        # Regex estricto para encontrar {{ variables }} y bucles
        vars_simples = re.findall(r'{{\s*(\w+)\s*}}', texto_completo)
        vars_bucles_tr = re.findall(r'{%\s*tr\s+for\s+\w+\s+in\s+(\w+)\s*%}', texto_completo)
        vars_bucles_p = re.findall(r'{%\s*p\s+for\s+\w+\s+in\s+(\w+)\s*%}', texto_completo)
        vars_bucles_simple = re.findall(r'{%\s*for\s+\w+\s+in\s+(\w+)\s*%}', texto_completo)
        
        # Unir todos los detectados en bucles para marcarlos como dinamicos
        todas_las_listas = set(vars_bucles_tr + vars_bucles_p + vars_bucles_simple)
        
        # Limpiar duplicados totales
        variables_detectadas = set(vars_simples + vars_bucles_tr + vars_bucles_p + vars_bucles_simple)
        
        esquema = []
        for var in variables_detectadas:
            if var.startswith("img_"):
                esquema.append({"nombre": var, "tipo": "imagen"})
            elif var in todas_las_listas or var.startswith("tbl_") or var.startswith("lista_"):
                esquema.append({"nombre": var, "tipo": "tabla_dinamica"})
            else:
                esquema.append({"nombre": var, "tipo": "texto"})
                
        return {"variables": esquema}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analizando documento: {str(e)}")"""

content = re.sub(r'@router.get\("/plantillas/\{id\}/esquema"\).*?raise HTTPException\(status_code=500, detail=f"Error procesando documento: \{str\(e\)\}"\)', new_func, content, flags=re.DOTALL)
content = re.sub(r'@router.get\("/plantillas/\{id\}/esquema"\).*?raise HTTPException\(500, f"Error analizando plantilla con jinja2: \{str\(e\)\}"\)', new_func, content, flags=re.DOTALL)

with open(path, "w") as f:
    f.write(content)
