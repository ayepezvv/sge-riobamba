from docx import Document

document = Document()
document.add_heading('Plantilla de TDR', 0)
document.add_paragraph('Proyecto: {{ nombre_proyecto }}')
document.add_paragraph('Fecha: {{ fecha }}')
document.save('templates/contratacion/dummy_template.docx')
